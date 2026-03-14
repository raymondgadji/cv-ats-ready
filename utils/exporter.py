"""
utils/exporter.py
US-10 : Export des résultats en PDF ou DOCX, en mémoire (bytes).
US-14 : Aucun fichier n'est écrit sur disque.
"""

import io
from datetime import date


# ─────────────────────────────────────────
# EXPORT PDF
# ─────────────────────────────────────────

def export_to_pdf(cv_text: str, cover_letter_text: str = "") -> bytes:
    """
    Génère un PDF en mémoire avec le CV et optionnellement la lettre de motivation.
    Utilise reportlab.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, PageBreak
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    ORANGE = colors.HexColor("#FF6B00")

    style_title = ParagraphStyle(
        "CVTitle",
        parent=styles["Normal"],
        fontSize=18,
        fontName="Helvetica-Bold",
        textColor=ORANGE,
        spaceAfter=6,
    )
    style_section = ParagraphStyle(
        "Section",
        parent=styles["Normal"],
        fontSize=11,
        fontName="Helvetica-Bold",
        textColor=ORANGE,
        spaceBefore=14,
        spaceAfter=4,
    )
    style_body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica",
        leading=15,
        spaceAfter=3,
    )
    style_watermark = ParagraphStyle(
        "Watermark",
        parent=styles["Normal"],
        fontSize=8,
        textColor=colors.HexColor("#BBBBBB"),
        alignment=TA_CENTER,
    )

    story = []

    # ── CV optimisé ──
    story.append(Paragraph("CV optimisé — ATS Ready", style_title))
    story.append(HRFlowable(width="100%", thickness=2, color=ORANGE, spaceAfter=12))

    for line in cv_text.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        # Détecte les titres de section (ligne en majuscules ou avec ═══)
        if line.isupper() and len(line) > 3:
            story.append(Paragraph(line, style_section))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#EEEEEE"), spaceAfter=4))
        else:
            # Échapper les caractères HTML spéciaux
            safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(safe_line, style_body))

    # Watermark bas de page
    story.append(Spacer(1, 20))
    story.append(Paragraph(f"Généré par cv-ats-ready.fr — {date.today().strftime('%d/%m/%Y')}", style_watermark))

    # ── Lettre de motivation (page 2 si présente) ──
    if cover_letter_text.strip():
        story.append(PageBreak())
        story.append(Paragraph("Lettre de motivation", style_title))
        story.append(HRFlowable(width="100%", thickness=2, color=ORANGE, spaceAfter=12))

        for line in cover_letter_text.split("\n"):
            line = line.strip()
            if not line:
                story.append(Spacer(1, 6))
            else:
                safe_line = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(Paragraph(safe_line, style_body))

        story.append(Spacer(1, 20))
        story.append(Paragraph(f"Généré par cv-ats-ready.fr — {date.today().strftime('%d/%m/%Y')}", style_watermark))

    doc.build(story)
    return buffer.getvalue()


# ─────────────────────────────────────────
# EXPORT DOCX
# ─────────────────────────────────────────

def export_to_docx(cv_text: str, cover_letter_text: str = "") -> bytes:
    """
    Génère un .docx en mémoire avec le CV et optionnellement la lettre de motivation.
    Utilise python-docx.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ORANGE_RGB = RGBColor(0xFF, 0x6B, 0x00)

    doc = Document()

    # ── Marges ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def add_cv_content(text: str, title: str):
        # Titre principal
        h = doc.add_heading(title, level=1)
        h.runs[0].font.color.rgb = ORANGE_RGB
        h.runs[0].font.size = Pt(16)

        # Ligne de séparation via paragraphe vide avec bordure
        sep = doc.add_paragraph()
        sep.paragraph_format.space_after = Pt(6)

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            # Titre de section détecté
            if line.isupper() and len(line) > 3:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                run = p.add_run(line)
                run.bold = True
                run.font.color.rgb = ORANGE_RGB
                run.font.size = Pt(11)
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(10)

        # Watermark footer
        footer_p = doc.add_paragraph(f"Généré par cv-ats-ready.fr — {date.today().strftime('%d/%m/%Y')}")
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # ── CV optimisé ──
    add_cv_content(cv_text, "CV optimisé — ATS Ready")

    # ── Lettre de motivation (section 2) ──
    if cover_letter_text.strip():
        doc.add_page_break()
        add_cv_content(cover_letter_text, "Lettre de motivation")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
